import subprocess
import sys
import json
import os
import re
import time
from typing import List, Dict, Tuple, Optional
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import hashlib
import tkinter as tk
from tkinter import filedialog
import traceback

class DeepSeekClient:
    """DeepSeek API客户端"""
    
    def __init__(self, api_key: str, base_url: str = "https://api.deepseek.com"):
        self.api_key = api_key
        self.base_url = base_url
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
    
    def call_api(self, messages: List[Dict], model: str = "deepseek-chat", max_retries: int = 3) -> str:
        """调用DeepSeek API，支持重试"""
        
        for attempt in range(max_retries):
            try:
                print(f"正在调用DeepSeek API (第{attempt+1}次尝试)...")
                
                data = {
                    "model": model,
                    "messages": messages,
                    "temperature": 0.7,
                    "max_tokens": 4000
                }
                
                response = requests.post(
                    f"{self.base_url}/chat/completions",
                    headers=self.headers,
                    json=data,
                    timeout=60
                )
                
                if response.status_code == 200:
                    result = response.json()
                    return result["choices"][0]["message"]["content"]
                elif response.status_code == 401:
                    raise Exception("API密钥错误，请检查config.json配置")
                elif response.status_code == 429:
                    wait_time = 5
                    print(f"API频率限制，等待{wait_time}秒后重试...")
                    time.sleep(wait_time)
                    continue
                else:
                    print(f"API返回错误状态码: {response.status_code}")
                    print(f"响应内容: {response.text[:200]}")
                    raise Exception(f"API调用失败: HTTP {response.status_code}")
                    
            except requests.exceptions.Timeout:
                print(f"API调用超时 (尝试 {attempt+1}/{max_retries})")
                if attempt < max_retries - 1:
                    wait_time = 2 * (attempt + 1)
                    print(f"等待{wait_time}秒后重试...")
                    time.sleep(wait_time)
                else:
                    raise Exception(f"API调用超时，已重试{max_retries}次")
                    
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"API调用失败，准备重试: {e}")
                    time.sleep(2)
                else:
                    raise Exception(f"API调用失败，已重试{max_retries}次: {e}")
        
        raise Exception(f"API调用失败，已重试{max_retries}次")

class ChinesePoemQuizGenerator:
    """高中语文情境式默写题目生成器"""
    
    def __init__(self, api_key: str):
        self.client = DeepSeekClient(api_key)
        self.poems = []
        self.answers = []  # 存储答案（包含完整诗句的情境化表述）
        self.questions = []  # 存储题目（答案的挖空版）
        self.knowledge_base = "knowledge_base.json"
    
    def select_file_via_gui(self) -> Optional[str]:
        """通过GUI对话框选择文件"""
        try:
            root = tk.Tk()
            root.withdraw()
            
            file_path = filedialog.askopenfilename(
                title="选择包含诗文的TXT文件",
                filetypes=[
                    ("文本文件", "*.txt"),
                    ("所有文件", "*.*")
                ],
                initialdir=os.path.expanduser("~/Desktop")
            )
            
            root.destroy()
            
            if file_path and os.path.exists(file_path):
                return file_path
            else:
                return None
                
        except Exception as e:
            print(f"GUI文件选择失败: {e}")
            return None
    
    def load_poems_from_txt(self, file_path: str) -> bool:
        """从txt文件加载诗文内容，简化版，只提取原文"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content.strip():
                print("错误：文件内容为空！")
                return False
            
            print(f"成功读取文件，内容长度：{len(content)} 字符")
            
            # 分析篇章结构
            analysis_prompt = f'''你是一位高中语文教学专家。请分析以下文本，识别出其中包含的部编版高中语文必备篇目。

文本内容：
{content}

请严格按照以下JSON格式返回分析结果，只提供篇目和完整原文：
{{
    "篇章数量": 整数,
    "篇章列表": [
        {{
            "标题": "篇目标题",
            "作者": "作者",
            "原文": "完整原文"
        }},
        ...
    ]
}}'''

            result = self.client.call_api([
                {"role": "system", "content": "你是一位专业的语文教学助手，擅长分析古典诗文。请严格遵守JSON格式要求。"},
                {"role": "user", "content": analysis_prompt}
            ])
            
            # 提取JSON
            json_match = re.search(r'\{.*\}', result, re.DOTALL)
            if not json_match:
                print("错误：未找到JSON格式的响应")
                return False
            
            json_str = json_match.group()
            data = json.loads(json_str)
            
            self.poems = data.get("篇章列表", [])
            if not self.poems:
                print("错误：未识别到任何篇章")
                return False
            
            print(f"成功识别出 {len(self.poems)} 个篇章")
            
            for poem in self.poems:
                print(f"\n处理篇目: {poem.get('标题')}")
                print(f"作者: {poem.get('作者')}")
                print(f"原文预览: {poem.get('原文', '')[:100]}...")
            
            return True
            
        except Exception as e:
            print(f"加载文件失败: {e}")
            traceback.print_exc()
            return False
    
    def generate_answers(self, num_answers: int) -> bool:
        """生成情境式默写答案（包含完整诗句的情境化表述）"""
        try:
            # 收集所有诗文原文
            all_poems_text = []
            for poem in self.poems:
                all_poems_text.append({
                    "标题": poem.get("标题", ""),
                    "作者": poem.get("作者", ""),
                    "原文": poem.get("原文", "")
                })
            
            if not all_poems_text:
                print("错误：没有可用的诗文内容")
                return False
            
            print(f"有 {len(all_poems_text)} 篇诗文可用")
            
            # 生成答案的提示词
            generation_prompt = f'''你是一位有20年教龄的高中语文特级教师，正在为高三学生编制情境式默写题目。

以下是可用的诗文原文：
{json.dumps(all_poems_text, ensure_ascii=False, indent=2)}

**请按照以下要求生成 {num_answers} 个情境式默写题目和答案：**

**质量要求**：
1. 情境真实：参考实际教学案例、高考真题、网络优质内容
2. 表述自然：避免生硬的"在...上"句式
3. 引用准确：确保诗句完整、标点正确
4. 语境合理：情境与诗句内涵高度契合
5. 多样性：每个答案应有不同的情境和角度

**格式要求 - 这是最重要的部分**：
1. 每个答案是一个完整的情境化表述
2. 每个情境化表述中必须包含且只包含两个需要默写的独立短句
3. 每个短句必须用**双中括号标记**，格式为：[[诗句内容]]
4. 两个短句之间用逗号分隔
5. 在返回的JSON中，除了context字段，还需要包含sentences字段，其中包含提取出的两个诗句
6. **必须包含is_continuous字段**：true表示两个诗句是连续的上下句，false表示不是

**关于独立短句的关键定义**：
- **句子分隔原则**：
  - 逗号（，）**一定是**分隔符，看到逗号必须分割
  - 顿号（、）**不是**分隔符，顿号连接的词语属于同一个短句
  - 句号（。）、分号（；）、问号（？）、感叹号（！）、冒号（：）也**是**分隔符
- 例如："学不可以已" 是一个短句
- 例如："青，取之于蓝" 应该分为两个短句：[[青]] 和 [[取之于蓝]]（逗号分割）
- 例如："木、金、水" 应该是一个短句：[[木、金、水]]（顿号不分割）
- 每个独立短句通常包含1-8个汉字，不包含标点符号

**关于连续性的重要说明**：
- 如果两个诗句在原文中是连续的上下句（用逗号、分号等连接），则is_continuous应为true
- 例如："青" 和 "取之于蓝" 是连续的上下句，is_continuous应为true
- 例如："学不可以已" 和 "青" 不是连续的，应该在各自独立的引号内

**重要注意事项**：
1. 诗句长度应该适中，通常不超过8个汉字
2. **单个字也算作一个句子**：如"青"、"学"等
3. 确保诗句是从原文中准确提取的
4. 不要修改诗句的原文内容
5. 确保标点符号正确
6. **严格遵守句子分隔原则**：看到逗号必须分割，顿号不分割
7. **特殊处理三连句**：对于"青，取之于蓝，而青于蓝"这样的三连句，应该避免使用，或者选择其中连续的两个短句

**完全正确的示例**：
答案1（连续句）：知乎上有人问："如何用古文论证学习能改变人的本性？"高赞回答引用了荀子的[[木直中绳]]，[[輮以为轮]]，说明外力改造的重要性。
- sentences: ["木直中绳", "輮以为轮"]
- is_continuous: true

答案2（单个字+连续句）：在课堂上，老师问："荀子如何描述学习的开始？"学生答道：[[青]]，[[取之于蓝]]，说明学习从基础开始。
- sentences: ["青", "取之于蓝"]
- is_continuous: true

**错误的示例**：
错误示例1：在课堂上，老师问："荀子认为学习能改变人的本性，他用了哪些比喻？"学生答道：[[木直中绳，輮以为轮]]，说明木材经过加工可以变弯。
（错误：将两个用逗号分隔的短句合并到一个标记中，应该分开标记）

错误示例2：在讲解比喻论证时，老师引用荀子的话：[[木、金、水]]，说明不同材料的特性。
（错误：顿号不是分隔符，不需要标记。但如果是"木、金"这样的列举，应该整体作为一个短句）

**JSON格式要求**：
请确保返回的JSON格式完全正确，所有字符串都必须用双引号引起，字符串内部的引号需要正确转义。

请生成 {num_answers} 个这样的答案，并以JSON格式返回：
{{
    "answers": [
        {{
            "id": 1,
            "context": "情境化表述，其中必须明确包含需要默写的两个独立短句，每个短句用[[ ]]标记，用逗号分隔",
            "sentences": ["第一个诗句", "第二个诗句"],
            "is_continuous": true
        }},
        ...
    ]
}}'''

            print(f"正在生成 {num_answers} 个答案...")
            result = self.client.call_api([
                {"role": "system", "content": "你是专业的高中语文特级教师，有丰富的教学经验，擅长编制高质量的高考模拟题。请严格按照要求，确保每个被引用的短句是独立的。特别注意：逗号（，）一定是分隔符，看到逗号必须分割；顿号（、）不是分隔符，顿号连接的词语属于同一个短句。每个短句单独用双中括号（[[ ]]）标记，不要合并多个短句。每个答案必须且只能包含两个被引用的短句，绝对不能超过两个。请特别关注连续性判断：如果两个诗句是连续的上下句，它们应该被标记为连续的。请确保返回的JSON格式完全正确。"},
                {"role": "user", "content": generation_prompt}
            ])
            
            print("API返回结果预览:", result[:200], "...")
            
            # 提取JSON - 增强错误处理和修复
            try:
                # 尝试直接解析
                data = json.loads(result)
            except json.JSONDecodeError as e:
                print(f"JSON解析失败: {e}")
                print("尝试修复JSON格式...")
                
                # 尝试修复JSON格式
                fixed_result = self._fix_json_format(result)
                
                # 再次尝试解析
                try:
                    data = json.loads(fixed_result)
                    print("JSON修复成功，继续处理...")
                except json.JSONDecodeError as e2:
                    print(f"JSON修复后仍然解析失败: {e2}")
                    print("尝试提取JSON对象...")
                    
                    # 尝试从文本中提取JSON对象
                    json_match = re.search(r'\{\s*"answers"\s*:\s*\[.*\]\s*\}', result, re.DOTALL)
                    if json_match:
                        try:
                            data = json.loads(json_match.group())
                            print("成功从文本中提取JSON对象")
                        except json.JSONDecodeError as e3:
                            print(f"提取的JSON解析失败: {e3}")
                            return False
                    else:
                        print("未找到JSON格式的内容")
                        return False
            
            if "answers" not in data:
                print("错误：响应中缺少'answers'字段")
                return False
            
            raw_answers = data["answers"]
            print(f"API返回了 {len(raw_answers)} 个原始答案")
            
            # 验证和修正答案
            self.answers = self._validate_and_fix_answers_api(raw_answers)
            
            print(f"成功生成 {len(self.answers)} 个答案")
            
            # 检查前3个答案
            if self.answers:
                print("检查前3个答案：")
                for i, answer in enumerate(self.answers[:3], 1):
                    context = answer.get('context', '')
                    sentences = answer.get('sentences', [])
                    is_continuous = answer.get('is_continuous', False)
                    print(f"答案{i}: 包含{len(sentences)}个短句，连续: {is_continuous}")
                    print(f"  内容: {context}")
                    print(f"  短句: {sentences}")
            else:
                print("警告：没有生成任何有效的答案！")
                print("原始答案信息：")
                for i, ans in enumerate(raw_answers[:3], 1):
                    print(f"原始答案{i}: {ans}")
            
            return True
            
        except Exception as e:
            print(f"生成答案失败: {e}")
            traceback.print_exc()
            return False
    
    def _fix_json_format(self, text: str) -> str:
        """修复JSON格式，特别是未转义的双引号"""
        if not text:
            return text
        
        # 简化修复：将context字段值内部的双引号转义
        lines = text.split('\n')
        result_lines = []
        
        for line in lines:
            if '"context"' in line and ':' in line:
                # 找到值的开始位置
                value_start = line.find(':"') + 2
                if value_start > 1:  # 找到了
                    # 提取值
                    remaining_line = line[value_start:]
                    
                    # 我们需要找到值的结束位置
                    value_end = -1
                    escape_next = False
                    
                    for j, char in enumerate(remaining_line):
                        if escape_next:
                            escape_next = False
                            continue
                        elif char == '\\':
                            escape_next = True
                        elif char == '"':
                            # 检查这个双引号是否是字符串的结束
                            if j + 1 < len(remaining_line):
                                next_char = remaining_line[j + 1]
                                if next_char in [',', '}', ']', ' ', '\t', '\n']:
                                    # 检查前一个字符是否不是转义字符
                                    if j > 0 and remaining_line[j - 1] != '\\':
                                        value_end = j
                                        break
                    
                    if value_end == -1:
                        # 如果没有找到合适的结束位置，使用简单方法
                        if remaining_line.endswith('",'):
                            value = remaining_line[:-2]
                        elif remaining_line.endswith('"'):
                            value = remaining_line[:-1]
                        else:
                            value = remaining_line
                    else:
                        value = remaining_line[:value_end]
                    
                    # 转义值内部的双引号
                    value = value.replace('"', '\\"')
                    
                    # 重新构建行
                    if value_end == -1:
                        if remaining_line.endswith('",'):
                            fixed_line = line[:value_start] + value + '",'
                        elif remaining_line.endswith('"'):
                            fixed_line = line[:value_start] + value + '"'
                        else:
                            fixed_line = line[:value_start] + value + '"'
                    else:
                        fixed_line = line[:value_start] + value + remaining_line[value_end:]
                    
                    result_lines.append(fixed_line)
                else:
                    result_lines.append(line)
            else:
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def _validate_and_fix_answers_api(self, raw_answers: List[Dict]) -> List[Dict]:
        """验证和修正API返回的答案"""
        valid_answers = []
        
        for i, answer in enumerate(raw_answers, 1):
            context = answer.get('context', '')
            if not context:
                print(f"警告：答案{i}没有context，跳过")
                continue
            
            # 提取sentences字段
            sentences = answer.get('sentences', [])
            if not sentences or len(sentences) < 2:
                print(f"警告：答案{i}没有足够的句子，跳过")
                continue
            
            # 获取is_continuous字段
            is_continuous = answer.get('is_continuous', False)
            
            # 验证context中是否包含对应的[[标记]]
            sentence1 = sentences[0] if len(sentences) > 0 else ""
            sentence2 = sentences[1] if len(sentences) > 1 else ""
            
            has_mark1 = f'[[{sentence1}]]' in context
            has_mark2 = f'[[{sentence2}]]' in context
            
            if not (has_mark1 and has_mark2):
                print(f"警告：答案{i}的context中没有包含所有诗句标记，跳过")
                continue
            
            # 关键修改：如果句子包含逗号，强制按逗号分割
            all_sentences = []
            
            for sentence in [sentence1, sentence2]:
                # 检查句子是否包含逗号
                if '，' in sentence or ',' in sentence:
                    # 按逗号分割
                    parts = re.split(r'[，,]', sentence)
                    # 过滤空字符串
                    parts = [p.strip() for p in parts if p.strip()]
                    all_sentences.extend(parts)
                else:
                    all_sentences.append(sentence)
            
            # 过滤空字符串
            all_sentences = [s.strip() for s in all_sentences if s.strip()]
            
            if len(all_sentences) >= 2:
                # 只取前两个句子
                final_sentences = all_sentences[:2]
                
                # 检查这两个句子在原文中是否是连续的
                is_continuous_final = self._check_if_continuous(final_sentences[0], final_sentences[1])
                
                # 符合要求，添加到有效答案
                valid_answers.append({
                    "id": i,
                    "context": context,
                    "sentences": final_sentences,
                    "is_continuous": is_continuous_final
                })
                print(f"  √ 答案{i}有效，最终句子: {final_sentences}，连续: {is_continuous_final}")
            else:
                print(f"警告：答案{i}处理后句子不足2个，跳过")
        
        return valid_answers
    
    def _check_if_continuous(self, sentence1: str, sentence2: str) -> bool:
        """检查两个句子是否是上下句（在原句中用逗号连接）"""
        # 清理句子的标点
        clean_s1 = sentence1.strip().rstrip('，。！？；、').rstrip('，')
        clean_s2 = sentence2.strip().rstrip('，。！？；、').rstrip('，')
        
        # 遍历所有诗篇
        for poem in self.poems:
            poem_text = poem.get("原文", "")
            
            # 检查两种可能的连接方式
            patterns = [
                f"{clean_s1}，{clean_s2}",  # 中文逗号连接
                f"{clean_s1}, {clean_s2}",   # 英文逗号+空格连接
                f"{clean_s1},{clean_s2}",    # 英文逗号连接
            ]
            
            for pattern in patterns:
                if pattern in poem_text:
                    return True
        
        return False
    
    def save_answers_to_txt(self, file_path: str = "answers_初版.txt") -> bool:
        """保存答案初版到txt文件"""
        try:
            if not self.answers:
                print("错误：没有可保存的答案")
                return False
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("情境式默写题目答案（初版）\n")
                f.write("=" * 50 + "\n\n")
                
                for i, answer in enumerate(self.answers, 1):
                    answer_id = answer.get('id', i)
                    context = answer.get('context', '')
                    
                    f.write(f"答案 {answer_id}：\n")
                    f.write(f"{context}\n")
                    f.write("-" * 40 + "\n\n")
            
            print(f"答案已保存到 {file_path}")
            print(f"共保存了 {len(self.answers)} 个答案")
            return True
            
        except Exception as e:
            print(f"保存答案文件失败: {e}")
            return False
    
    def generate_questions_from_answers(self):
        """从答案生成题目（将[[标记]]替换为横线）"""
        if not self.answers:
            print("错误：没有可生成题目的答案")
            return []
        
        self.questions = []
        print(f"【调试】开始生成题目，共有 {len(self.answers)} 个答案需要处理。")
        
        for i, answer in enumerate(self.answers, 1):
            context = answer.get('context', '')
            sentences = answer.get('sentences', [])
            is_continuous = answer.get('is_continuous', False)
            
            if not context:
                print(f"警告：答案{i}内容为空，跳过")
                continue
            
            print(f"\n--- 处理第 {i} 题 ---")
            print(f"原始答案: {context}")
            print(f"提取到的诗句: {sentences}")
            print(f"是否连续: {is_continuous}")
            
            # 创建题目：将[[标记]]替换为横线
            question_text = context
            
            if len(sentences) >= 2:
                sentence1 = sentences[0]
                sentence2 = sentences[1]
                
                # 根据是否是上下句使用不同的格式
                if is_continuous:
                    # 上下句：合并为一个"__________，__________"格式
                    # 尝试匹配连续句模式
                    continuous_pattern = f'[[{sentence1}]]，[[{sentence2}]]'
                    if continuous_pattern in question_text:
                        question_text = question_text.replace(continuous_pattern, '"__________，__________"')
                        print(f"  替换连续句模式: {continuous_pattern} -> '__________，__________'")
                    else:
                        # 如果不是连续模式，分别替换
                        question_text = question_text.replace(f'[[{sentence1}]]', '"__________"')
                        question_text = question_text.replace(f'[[{sentence2}]]', '"__________"')
                        print(f"  分别替换非连续句")
                else:
                    # 非上下句：使用"__________"，"__________"格式
                    question_text = question_text.replace(f'[[{sentence1}]]', '"__________"')
                    question_text = question_text.replace(f'[[{sentence2}]]', '"__________"')
                    print(f"  分别替换非连续句")
            
            print(f"最终生成的题目: {question_text}")
            
            # 存储题目
            self.questions.append({
                "id": i,
                "question": f"{i}. {question_text}",
                "answer": context,
                "sentences": sentences[:2],
                "is_continuous": is_continuous
            })
        
        print(f"\n{'='*60}")
        print(f"题目生成完毕！成功生成了 {len(self.questions)} 道题目。")
        if self.questions:
            print("前3题预览:")
            for q in self.questions[:3]:
                print(f"  Q{q['id']}: {q['question']}")
        print(f"{'='*60}")
        
        return self.questions
    
    def save_to_docx(self) -> Optional[str]:
        """保存为Word文档，题目黑体三号，正文黑色五号"""
        if not self.questions:
            print("错误：没有可保存的题目")
            return None
        
        try:
            doc = Document()
            
            # 设置标题
            title = doc.add_heading('高中语文情境式默写试题', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置标题字体为黑体三号
            for run in title.runs:
                run.font.size = Pt(16)  # 三号字约16磅
                run.font.name = '黑体'
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), '黑体')
            
            # 添加考试信息
            info_para = doc.add_paragraph()
            info_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
            
            for run in info_para.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 添加题目
            for q in self.questions:
                p = doc.add_paragraph(q['question'])
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(12)
                
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 添加分页符
            doc.add_page_break()
            
            # 添加答案标题
            answer_title = doc.add_heading('参考答案', 0)
            answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in answer_title.runs:
                run.font.size = Pt(16)
                run.font.name = '黑体'
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), '黑体')
            
            # 添加答案 - 将[[标记]]替换为中文引号
            for q in self.questions:
                answer_text = q['answer']
                sentences = q.get('sentences', [])
                is_continuous = q.get('is_continuous', False)
                
                # 将答案中的[[标记]]替换为中文引号
                if sentences and len(sentences) >= 2:
                    sentence1 = sentences[0]
                    sentence2 = sentences[1]
                    
                    # 替换第一个标记
                    answer_text = answer_text.replace(f'[[{sentence1}]]', f'"{sentence1}"')
                    # 替换第二个标记
                    answer_text = answer_text.replace(f'[[{sentence2}]]', f'"{sentence2}"')
                    
                    # 如果是上下句，合并为一个引号
                    if is_continuous:
                        # 查找模式："句子1"，"句子2"
                        pattern1 = f'"{sentence1}"，"{sentence2}"'
                        replacement1 = f'"{sentence1}，{sentence2}"'
                        answer_text = answer_text.replace(pattern1, replacement1)
                        
                        # 也处理其他可能的标点
                        pattern2 = f'"{sentence1}", "{sentence2}"'
                        replacement2 = f'"{sentence1}，{sentence2}"'
                        answer_text = answer_text.replace(pattern2, replacement2)
                
                p = doc.add_paragraph(f"{q['id']}. {answer_text}")
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(12)
                
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            
            # 保存文件到C盘根目录"情境默写生成结果"文件夹
            output_folder = "C:\\情境默写生成结果"
            os.makedirs(output_folder, exist_ok=True)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"情境式默写题目_{timestamp}.docx"
            full_path = os.path.join(output_folder, filename)
            doc.save(full_path)
            
            # 返回完整的文件路径和文件夹路径
            return {
                "file_path": full_path,
                "folder_path": output_folder
            }
            
        except Exception as e:
            print(f"保存Word文档失败: {e}")
            traceback.print_exc()
            return None
    
    def run_cli(self):
        """命令行界面运行程序（自动模式）"""
        print("="*50)
        print("高中语文情境式默写题目生成系统")
        print("（整合版，API处理所有文本预处理）")
        print("="*50)
        
        # 1. 读取txt文件 - 直接使用GUI选择
        print("\n请使用GUI对话框选择诗文文件...")
        txt_file = self.select_file_via_gui()
        
        if not txt_file or not os.path.exists(txt_file):
            print("文件选择取消或文件不存在，程序退出。")
            return
        
        print(f"\n正在分析诗文内容: {txt_file}")
        
        if not self.load_poems_from_txt(txt_file):
            print("加载诗文失败，请检查文件格式")
            return
        
        # 显示识别结果
        for i, poem in enumerate(self.poems, 1):
            print(f"{i}. {poem.get('标题', '未知标题')} - {poem.get('作者', '未知作者')}")
        
        # 2. 输入题目数量
        while True:
            try:
                num_input = input("\n请输入需要生成的题目数量: ").strip()
                if not num_input:
                    continue
                    
                num_answers = int(num_input)
                if num_answers > 0:
                    break
                print("请输入正整数")
            except ValueError:
                print("请输入有效的数字")
        
        # 3. 生成答案
        print(f"\n正在生成 {num_answers} 个情境式默写答案...")
        if not self.generate_answers(num_answers):
            print("答案生成失败")
            return
        
        # 4. 从答案生成题目
        print("\n正在从答案生成题目（将诗句替换为横线）...")
        self.generate_questions_from_answers()
        
        if not self.questions:
            print("生成题目失败")
            return
        
        # 5. 保存Word文档
        save_result = self.save_to_docx()
        
        print("\n" + "="*50)
        print("试卷生成完成！")
        
        if save_result and isinstance(save_result, dict):
            file_path = save_result.get("file_path", "")
            folder_path = save_result.get("folder_path", "")
            
            if file_path and folder_path:
                # 告诉用户输出文件夹位置
                print(f"Word文档已生成：{file_path}")
                print(f"\n输出文件夹位置：{folder_path}")
                
                # 尝试自动打开输出文件夹
                try:
                    if os.name == 'nt':  # Windows
                        os.startfile(folder_path)
                    elif os.name == 'posix':  # macOS/Linux
                        if sys.platform == 'darwin':  # macOS
                            subprocess.run(['open', folder_path])
                        else:  # Linux
                            subprocess.run(['xdg-open', folder_path])
                    print("已自动打开输出文件夹。")
                except Exception as e:
                    print(f"无法自动打开文件夹，请手动访问：{folder_path}")
                    print(f"错误信息：{e}")
                
                # 等待用户按回车键
                input("\n按回车键关闭程序...")
                return
            else:
                print("错误：无法获取保存的文件信息。")
        else:
            print("错误：Word文档保存失败。")
        
        print("="*50)
    
    def run(self, use_gui: bool = False):
        """主运行流程"""
        try:
            self.run_cli()
        except KeyboardInterrupt:
            print("\n程序被用户中断")
        except Exception as e:
            print(f"程序运行出错: {e}")
            traceback.print_exc()
            input("\n按Enter键退出...")

# 主程序入口
if __name__ == "__main__":
    print("高中语文情境式默写题目生成系统")
    print("="*50)
    
    # 获取可执行文件所在目录
    if getattr(sys, 'frozen', False):
        # 如果是PyInstaller打包后的exe
        application_path = os.path.dirname(sys.executable)
    else:
        # 如果是直接运行脚本
        application_path = os.path.dirname(os.path.abspath(__file__))
    
    # 配置文件的路径
    config_file = os.path.join(application_path, "config.json")
    
    if not os.path.exists(config_file):
        print("错误：找不到配置文件 config.json")
        print(f"配置文件应该放在: {application_path}")
        print("请创建 config.json 文件，内容如下：")
        print('''{
  "deepseek_api_key": "your_api_key_here",
  "knowledge_base_path": "knowledge_base.json",
  "default_output_dir": "./output"
}''')
        input("按Enter键退出...")
    else:
        # 加载配置
        try:
            with open(config_file, "r", encoding="utf-8") as f:
                config = json.load(f)
            
            api_key = config.get("deepseek_api_key", "")
            if not api_key or api_key == "your_api_key_here":
                print("错误：请先在 config.json 中配置你的DeepSeek API密钥")
                input("按Enter键退出...")
            else:
                # 运行程序
                generator = ChinesePoemQuizGenerator(api_key)
                generator.run()
                
        except json.JSONDecodeError:
            print("错误：配置文件格式不正确")
            input("按Enter键退出...")
        except Exception as e:
            print(f"程序启动失败: {e}")
            input("按Enter键退出...")